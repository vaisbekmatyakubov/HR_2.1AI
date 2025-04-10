import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- Инициализация состояния сессии ---
if 'analysis_done' not in st.session_state:
    st.session_state.analysis_done = False # Флаг: был ли выполнен анализ
if 'gemini_results_dict' not in st.session_state:
    st.session_state.gemini_results_dict = {} # Словарь для хранения результатов Gemini
if 'edited_results' not in st.session_state:
     st.session_state.edited_results = {} # Для хранения отредактированных значений (не используется в этой версии, читаем напрямую)


# Загрузка переменных окружения
load_dotenv()

# --- Конфигурация Gemini API ---
try:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        api_key = st.secrets["GEMINI_API_KEY"]

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash-latest')
except Exception as e:
    st.error(f"Хатолик: Gemini API калитини созлашда муаммо. Калит мавжудлигини текширинг. {e}")
    st.stop()

# --- Функции (extract_komandirovka_info, parse_gemini_output, create_docx_report) ---
# Они остаются такими же, как в предыдущем ответе.
# Важно: убедитесь, что у вас есть актуальные версии этих функций.
# Особенно важно: промпт в extract_komandirovka_info все еще рекомендуется упростить!

def extract_komandirovka_info(text):
    """Отправляет текст и промпт в Gemini."""
    # ВАШ СЛОЖНЫЙ ПРОМПТ ЗДЕСЬ (РЕКОМЕНДУЕТСЯ УПРОСТИТЬ)
    # extract_komandirovka_info функцияси ичидаги промптни ўзгартиринг:
    prompt = f"""
    Сен Ўзбек (кирилл) тилида ёзилган командировка (хизмат сафари) тавсифларидан аниқ маълумотларни ажратиб олиш учун мўлжалланган ёрдамчисан.

    Қуйидаги командировка тавсифидан керакли маълумотларни ажратиб ол:

    "{text}"

    Қуйидаги форматда, ҳар бир маълумотни янги қатордан бошлаб, жавоб бер:
    Исм: [топилган исм]
    Лавозим: [топилган лавозим]
    Бошланиш санаси: [топилган сана]
    Буйрук санаси: [кайси санада хизмат сафарига чкиш учун буйрик берилган]
    Тугаш санаси: [топилган сана]
    Жой: [топилган жой/шаҳар/манзил]
    Мақсад: [топилган мақсад]  # <--- СОДДАЛАШТИРИЛДИ!
    Кунлик Ҳисобот: [кунлик ҳисоботни йози бер озинг толик тасвирлаб бер] # (Буни ҳам соддалаштириш керак)
    Сафар Натижаси: [топилган сафар натижаси озинг толик тасвирлаб бер] # (Буни ҳам соддалаштириш керак)
    Хулоса: [топилган хулоса ёки натижани озинг толик тасвирлаб бер ] # (Буни ҳам соддалаштириш керак)
    Умумий ксим: [берилган промпдан келиб чиким 1чи шахис томонидан толикрок килиб тасвирлаб бер...] # (Буни ҳам соддалаштириш керак)

    Агар бирор маълумот матнда мавжуд бўлмаса ёки аниқ бўлмаса, "Топилмади" деб ёз.
    Жавобни фақат ўзбек тилида, юқоридаги форматда қайтар. Бошқа ҳеч қандай қўшимча матн ёзма.
    """
    # ... (функциянинг қолган қисми)


    try:
        response = model.generate_content(prompt)
        if response.parts:
            cleaned_response = response.text.strip().replace("```", "").strip()
            return cleaned_response
        else:
            try:
                feedback = response.prompt_feedback
                block_reason = feedback.block_reason if hasattr(feedback, 'block_reason') else 'Номаълум сабаб'
                return f"Хатолик: Gemini жавоб қайтармади. Блоклаш сабаби: {block_reason}"
            except Exception:
                return "Хатолик: Gemini дан бўш жавоб қайтди (эҳтимол хавфсизлик фильтри)."
    except Exception as e:
        st.error(f"Gemini API га мурожаат қилишда хатолик: {e}")
        return "Маълумот олишда хатолик юз берди."

def parse_gemini_output(gemini_text):
    """Разбирает многострочный ответ Gemini на словарь."""
    data = {}
    if not gemini_text or not isinstance(gemini_text, str):
        return data
    lines = gemini_text.strip().split('\n')
    for line in lines:
        parts = line.split(':', 1)
        if len(parts) == 2:
            key = parts[0].strip()
            value = parts[1].strip()
            data[key] = value
    return data

def create_docx_report(gemini_data, manual_data):
    """Создает DOCX документ в памяти."""
    document = Document()
    title = document.add_paragraph('ХИЗМАТ САФАРИ БЎЙИЧА МАЪЛУМОТЛАР')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(14)
        run.bold = True
    document.add_paragraph()

    if manual_data.get("report_date"):
        p = document.add_paragraph()
        p.add_run('Ҳисобот Санаси: ').bold = True
        if isinstance(manual_data["report_date"], datetime.date):
             p.add_run(manual_data["report_date"].strftime('%Y-%m-%d'))
        else:
             p.add_run(str(manual_data["report_date"]))
    if manual_data.get("dept_head_name"):
        p = document.add_paragraph()
        p.add_run('Бўлим бошлиғи Исми Фамилияси: ').bold = True
        p.add_run(manual_data["dept_head_name"])
    if manual_data.get("dept_head_position"):
        p = document.add_paragraph()
        p.add_run('Бўлим бошлиғи Лавозими: ').bold = True
        p.add_run(manual_data["dept_head_position"])
    document.add_paragraph()

    if gemini_data:
        p = document.add_paragraph()
        p.add_run('Хизмат сафари тафсилотлари:').bold = True # Убрали (AI)
        document.add_paragraph()
        for key, value in gemini_data.items():
            p = document.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(value)

    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# --- Интерфейс Streamlit ---
st.set_page_config(page_title="Командировка Таҳлили", layout="wide")

st.title("📄 ДОК")
st.markdown("Командировка (хизмат сафари) ҳақидаги матнни киритинг ва керакли маълумотларни ажратиб олинг.")

# --- Блок ввода дополнительных данных ---
st.divider()
st.subheader("Қўшимча Маълумотлар (Фойдаланувчи томонидан киритилади):")
# Используем уникальные ключи для виджетов, чтобы их состояние сохранялось
dept_head_name = st.text_input("Бўлим бошлиғи Исми Фамилияси:", key="dept_head_name_input")
dept_head_position = st.text_input("Бўлим бошлиғи Лавозими:", key="dept_head_position_input")
report_date = st.date_input("Ҳисобот Санаси:", value=datetime.date.today(), key="report_date_input")
st.divider()
# --- Конец блока ---

# Поле для ввода текста командировки
user_input = st.text_area("Командировка тавсифини шу ерга киритинг:", height=150, placeholder="Мисол: ...") # Уменьшил высоту

# Кнопка "Анализ қилиш"
if st.button("🔍 Таҳлил қилиш", key="analyze_button"):
    if user_input:
        with st.spinner('Gemini фикрлаяпти... Илтимос, кутинг...'):
            extracted_data_str = extract_komandirovka_info(user_input)
            # Сохраняем результат парсинга в состояние сессии
            st.session_state.gemini_results_dict = parse_gemini_output(extracted_data_str)
            # Устанавливаем флаг, что анализ выполнен
            st.session_state.analysis_done = True
    else:
        st.warning("Илтимос, таҳлил қилиш учун командировка тавсифини киритинг.")
        st.session_state.analysis_done = False # Сбрасываем флаг, если текста нет
        st.session_state.gemini_results_dict = {} # Очищаем старые результаты

# --- Блок отображения и редактирования результатов Gemini ---
# Этот блок показывается ТОЛЬКО ПОСЛЕ успешного анализа
if st.session_state.get('analysis_done', False):
    st.divider()
    st.subheader("Таҳрирлаш учун майдонлар (Gemini натижалари):")

    # Словарь для хранения ТЕКУЩИХ значений из полей редактирования
    current_edited_values = {}
    # Поля, которые могут быть длинными
    long_text_keys = ["Мақсад", "Сафар Натижаси", "Хулоса", "Умумий ксим"]

    # Динамически создаем поля для редактирования
    for key, value in st.session_state.gemini_results_dict.items():
        field_key = f"edit_{key}" # Уникальный ключ для каждого виджета
        if key in long_text_keys:
             # Используем text_area для потенциально длинных полей
             current_value = st.text_area(f"**{key}:**", value=value, key=field_key, height=100)
        else:
             # Используем text_input для остальных
             current_value = st.text_input(f"**{key}:**", value=value, key=field_key)
        # Сохраняем текущее значение из виджета
        current_edited_values[key] = current_value

    st.divider()

    # Кнопка "Подтвердить и Создать DOCX"
    if st.button("✅ Тасдиқлаш ва DOCX яратиш", key="confirm_button"):
        # Собираем данные, введенные вручную (читаем их текущее состояние)
        manual_data_dict = {
            "report_date": report_date, # Берется из виджета выше
            "dept_head_name": dept_head_name, # Берется из виджета выше
            "dept_head_position": dept_head_position # Берется из виджета выше
        }

        # Создаем DOCX файл в памяти с ТЕКУЩИМИ (отредактированными) значениями
        try:
            # Передаем словарь с текущими значениями из полей редактирования
            docx_buffer = create_docx_report(current_edited_values, manual_data_dict)

            # Отображаем кнопку скачивания
            st.download_button(
                label="📄 DOCX файлни юклаб олиш",
                data=docx_buffer,
                file_name="komandirovka_hisoboti.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_button" # Добавили ключ кнопке скачивания
            )
            st.success("DOCX файл тайёр!")
        except Exception as docx_e:
             st.error(f"DOCX файлини яратишда хатолик: {docx_e}")


st.markdown("---")
st.caption("Powered by Google Gemini & Vaisbek Matyakubov")